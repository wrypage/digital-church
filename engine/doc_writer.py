#!/usr/bin/env python3
"""
engine/doc_writer.py

Creates clean Google-style .docx output
for Elias climate reports.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def _add_heading(doc, text):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_subheading(doc, text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(12)


def _add_quote(doc, text):
    p = doc.add_paragraph(text)
    p.style = 'Intense Quote'


def write_doc(output_lines, output_dir="out"):
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Clean Google-style default
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = f"Digital Pulpit – Climate Report"
    _add_heading(doc, title)

    _add_paragraph(doc, f"Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    for line in output_lines:
        if line.startswith("##"):
            _add_subheading(doc, line.replace("##", "").strip())
        elif line.startswith(">"):
            _add_quote(doc, line.replace(">", "").strip())
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            _add_paragraph(doc, line)

    filename = os.path.join(
        output_dir,
        f"climate_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    doc.save(filename)

    return filename